"""
Concept note generator for the Go Vocal competitor-battlecard skill.

Reads a briefing JSON (see assets/example_briefing_decidim.json for schema),
produces a branded .docx that applies the Go Vocal identity (Chivo font,
dark purple / medium purple / cherry palette) and localises every fixed
label via scripts/localisation.py.

Usage
-----
    python generate_concept_note.py <briefing.json> <output_path.docx>

The briefing author is responsible for writing the body prose in the target
language; this script only localises the fixed headings, labels and legend.

Schema (summary)
----------------
{
  "competitor": "Decidim",
  "language": "en",
  "cover": {
    "label_left": "GO VOCAL",
    "label_right": "DECIDIM",
    "title": "Upgrading your\nparticipation tools.",
    "subtitle": "More inclusive, more interactive..."
  },
  "intro_paragraphs": [ "You already run...", "..." ],
  "numbers": [
    { "value": "4.7 / 5", "label": "Customer Satisfaction Score" },
    ...  # up to 6 items
  ],
  "section_intro": "Across the migration conversations we've had...",
  "usps": [
    {
      "title": "Hybrid input & representativeness",
      "opportunity": "A recurring theme...",
      "features_intro": "Three capabilities...",
      "features": [
        { "name": "360° Input Portal", "desc": "One workspace..." },
        ...
      ]
    },
    ... # exactly 5
  ],
  "bottom_line": {
    "paragraphs": [ "Decidim is a serious...", "..." ],
    "closer_italic": "For OSP-hosted Decidim clients..."
  }
}

Any field the briefing omits is rendered empty rather than blocking the run,
so the user can regenerate partial drafts during the human-validation step.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Make the localisation module importable whether this script is launched
# from within scripts/ or from the skill root.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from localisation import L  # noqa: E402


# ---------------------------------------------------------------------------
# Brand tokens
# ---------------------------------------------------------------------------
DARK_PURPLE = RGBColor(0x1E, 0x15, 0x5D)
MEDIUM_PURPLE = RGBColor(0x43, 0x36, 0x9B)
CHERRY = RGBColor(0xFF, 0x3E, 0x52)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LILAC_1 = "F0EEFA"  # hex without leading '#'
DARK_PURPLE_HEX = "1E155D"

FONT = "Chivo"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------
def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_run(run, *, size=11, bold=False, italic=False, color=DARK_PURPLE):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_body(doc, text, *, size=11, bold=False, italic=False,
             color=DARK_PURPLE, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(doc, text, level=1):
    """Simple heading styles (H1 / H2 / H3) in Go Vocal purple."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.color.rgb = DARK_PURPLE
    if level == 1:
        run.font.size = Pt(28)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        run.font.size = Pt(22)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    else:
        run.font.size = Pt(16)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, bold_lead, body):
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(bold_lead + ". ")
    set_run(r1, size=11, bold=True)
    r2 = p.add_run(body)
    set_run(r2, size=11)
    p.paragraph_format.space_after = Pt(4)
    return p


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------
def setup_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_PURPLE

    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    return doc


def render_cover(doc: Document, briefing: dict, lang: str) -> None:
    cover = briefing.get("cover", {})
    competitor = briefing.get("competitor", "")
    label_left = cover.get("label_left", "GO VOCAL")
    label_right = cover.get("label_right", competitor.upper())
    connector = L(lang, "cover_label_connector")

    # Top label
    p_label = doc.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_left = p_label.add_run(label_left)
    set_run(r_left, size=13, bold=True)
    r_x = p_label.add_run(connector)
    set_run(r_x, size=13, bold=True, color=CHERRY)
    r_right = p_label.add_run(label_right)
    set_run(r_right, size=13, bold=True)

    # Vertical spacing before the title
    for _ in range(3):
        doc.add_paragraph()

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_title.add_run(cover.get("title", ""))
    set_run(run, size=46)
    p_title.paragraph_format.line_spacing = 1.05
    p_title.paragraph_format.space_after = Pt(28)

    # Subtitle
    subtitle = cover.get("subtitle", "")
    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_sub.add_run(subtitle)
        set_run(run, size=13, italic=True, color=MEDIUM_PURPLE)
        p_sub.paragraph_format.line_spacing = 1.4

    page_break(doc)


def render_why_this_note(doc: Document, briefing: dict, lang: str) -> None:
    add_heading(doc, L(lang, "why_this_note"), level=2)
    for para in briefing.get("intro_paragraphs", []):
        add_body(doc, para, size=11)


def render_numbers(doc: Document, briefing: dict, lang: str) -> None:
    numbers = briefing.get("numbers", [])
    if not numbers:
        return
    add_heading(doc, L(lang, "go_vocal_in_numbers"), level=3)

    cols = len(numbers)
    table = doc.add_table(rows=2, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, item in enumerate(numbers):
        top = table.cell(0, i)
        bot = table.cell(1, i)

        set_cell_bg(top, DARK_PURPLE_HEX)
        top.text = ""
        p = top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(item.get("value", ""))
        set_run(r, size=20, bold=True, color=WHITE)
        top.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        set_cell_bg(bot, LILAC_1)
        bot.text = ""
        p = bot.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(item.get("label", ""))
        set_run(r, size=9)
        bot.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def render_usp_header(doc: Document, number: int, title: str) -> None:
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    t.columns[0].width = Cm(2.5)
    t.columns[1].width = Cm(14)
    t.rows[0].cells[0].width = Cm(2.5)
    t.rows[0].cells[1].width = Cm(14)

    num_cell = t.rows[0].cells[0]
    set_cell_bg(num_cell, DARK_PURPLE_HEX)
    num_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = num_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(number))
    set_run(r, size=32, bold=True, color=WHITE)

    title_cell = t.rows[0].cells[1]
    set_cell_bg(title_cell, LILAC_1)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = title_cell.paragraphs[0]
    r = p.add_run("  " + title)
    set_run(r, size=20)

    doc.add_paragraph()  # spacing after the header


def render_usp(doc: Document, number: int, usp: dict, lang: str) -> None:
    render_usp_header(doc, number, usp.get("title", ""))

    # Opportunity label (cherry)
    p = doc.add_paragraph()
    r = p.add_run(L(lang, "the_opportunity"))
    set_run(r, size=12, bold=True, color=CHERRY)
    p.paragraph_format.space_after = Pt(2)
    add_body(doc, usp.get("opportunity", ""), size=11)

    # What we offer label (cherry)
    p = doc.add_paragraph()
    r = p.add_run(L(lang, "what_we_offer"))
    set_run(r, size=12, bold=True, color=CHERRY)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)

    intro = usp.get("features_intro")
    if intro:
        add_body(doc, intro, size=11)

    for feat in usp.get("features", []):
        add_bullet(doc, feat.get("name", ""), feat.get("desc", ""))

    doc.add_paragraph()  # trailing spacing


def render_usps(doc: Document, briefing: dict, lang: str) -> None:
    page_break(doc)
    add_heading(doc, L(lang, "where_we_stand_out"), level=2)
    section_intro = briefing.get("section_intro")
    if section_intro:
        add_body(doc, section_intro, size=11)

    usps = briefing.get("usps", [])
    for i, usp in enumerate(usps, start=1):
        render_usp(doc, i, usp, lang)


def render_bottom_line(doc: Document, briefing: dict, lang: str) -> None:
    bottom = briefing.get("bottom_line") or {}
    if not bottom:
        return
    page_break(doc)
    add_heading(doc, L(lang, "the_bottom_line"), level=2)
    for para in bottom.get("paragraphs", []):
        add_body(doc, para, size=11)

    closer = bottom.get("closer_italic")
    if closer:
        add_body(doc, closer, size=11, italic=True, color=MEDIUM_PURPLE)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
def build(briefing: dict, output_path: Path) -> None:
    lang = briefing.get("language", "en")
    doc = setup_document()
    render_cover(doc, briefing, lang)
    render_why_this_note(doc, briefing, lang)
    render_numbers(doc, briefing, lang)
    render_usps(doc, briefing, lang)
    render_bottom_line(doc, briefing, lang)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        sys.stderr.write(
            "Usage: python generate_concept_note.py <briefing.json> <output_path.docx>\n"
        )
        return 2

    briefing_path = Path(argv[1])
    output_path = Path(argv[2])

    with briefing_path.open("r", encoding="utf-8") as f:
        briefing = json.load(f)

    build(briefing, output_path)
    print(f"Saved: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
