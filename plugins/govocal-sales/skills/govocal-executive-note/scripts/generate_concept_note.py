"""
Concept note generator for the Go Vocal concept-note skill.

Reads a briefing JSON and produces a branded .docx applying the Go Vocal
identity. Supports two modes via the top-level "mode" field:

  - "greenfield"  (default for city concept notes): single partnership cover
    label, optional logo on top, optional cover hero image, optional one image
    per USP. No competitor, no spreadsheet.
  - "competitive" (legacy battlecard): "GO VOCAL x COMPETITOR" cover label.
    Pairs with generate_spreadsheet.py. Backward compatible — briefings that
    omit "mode" and omit image specs render exactly as before.

Brand fonts (per govocal-brand): Chivo for headings/titles/numbers,
Libre Franklin for all running body copy. Both are Google Fonts and must be
installed on the machine opening the .docx for correct rendering; otherwise
the viewer falls back to a system sans-serif.

Image specs are { "path": <file or null>, "caption": <str>, "width_cm"?, "placement"? }.
A USP image renders centred BELOW that section's text when the file resolves
(`placement: "full"` → 16cm for wide shots, else 13cm). If the file is missing,
nothing is drawn — no placeholder box.

Usage
-----
    python generate_concept_note.py <briefing.json> <output.docx> [assets_dir]
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from localisation import L  # noqa: E402

# ---------------------------------------------------------------------------
# Brand tokens
# ---------------------------------------------------------------------------
DARK_PURPLE = RGBColor(0x1E, 0x15, 0x5D)
MEDIUM_PURPLE = RGBColor(0x43, 0x36, 0x9B)
CHERRY = RGBColor(0xFF, 0x3E, 0x52)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LILAC_1 = "F0EEFA"
DARK_PURPLE_HEX = "1E155D"

FONT_HEADING = "Chivo"          # headings, cover title, USP titles, numbers
FONT_BODY = "Libre Franklin"    # all running body copy (brand rule)

# Module-level so relative image paths in the briefing can be resolved.
ASSETS_DIR = Path(".")

CENTER = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------
def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def set_run(run, *, size=11, bold=False, italic=False, color=DARK_PURPLE,
            font=FONT_BODY):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_body(container, text, *, size=11, bold=False, italic=False,
             color=DARK_PURPLE, align=None, space_after=6, font=FONT_BODY):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, bold=bold, italic=italic,
            color=color, font=font)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(doc, text, level=1):
    """Heading styles (H1/H2/H3) — Chivo, Go Vocal purple."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_HEADING
    run.font.color.rgb = DARK_PURPLE
    if level == 1:
        run.font.size = Pt(28)
        p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(18), Pt(12)
    elif level == 2:
        run.font.size = Pt(22)
        p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(16), Pt(8)
    else:
        run.font.size = Pt(16)
        run.font.bold = True
        p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(12), Pt(6)
    return p


def add_bullet(container, bold_lead, body):
    p = container.add_paragraph(style="List Bullet")
    set_run(p.add_run(bold_lead + ". "), size=11, bold=True)   # Libre Franklin
    set_run(p.add_run(body), size=11)
    p.paragraph_format.space_after = Pt(4)
    return p


def _resolve(path: str):
    if not path:
        return None
    p = Path(path)
    for cand in (p, ASSETS_DIR / p):
        if cand.exists():
            return cand
    return None


def add_image_slot(doc, spec, width_cm):
    """Render an image (centred, below text) only if the file resolves.

    Returns True if an image was placed, False otherwise. No placeholder box is
    ever drawn — a missing image simply renders nothing.
    """
    if not spec:
        return False
    resolved = _resolve(spec.get("path"))
    if not resolved:
        return False
    width_cm = spec.get("width_cm", width_cm)
    doc.add_picture(str(resolved), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = CENTER
    caption = spec.get("caption", "")
    if caption:
        pc = doc.add_paragraph()
        pc.alignment = CENTER
        set_run(pc.add_run(caption), size=8, italic=True, color=MEDIUM_PURPLE)
    return True


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------
def setup_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_PURPLE
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = section.bottom_margin = Cm(2.2)
        section.left_margin = section.right_margin = Cm(2.2)
    return doc


# ---------------------------------------------------------------------------
# Cover (mode-aware)
# ---------------------------------------------------------------------------
def render_cover(doc: Document, briefing: dict, lang: str) -> None:
    cover = briefing.get("cover", {})
    mode = briefing.get("mode", "competitive")

    # Logos on top: Go Vocal brand logo, then the city's own logo / seal.
    logo_rendered = add_image_slot(doc, cover.get("logo"), 4.0)
    if logo_rendered:
        doc.add_paragraph()
    city_rendered = add_image_slot(doc, cover.get("city_logo") or cover.get("hero"), 5.5)
    any_logo = logo_rendered or city_rendered
    if any_logo:
        doc.add_paragraph()

    # Label line
    p_label = doc.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if mode == "competitive":
        label_left = cover.get("label_left", "GO VOCAL")
        label_right = cover.get("label_right", briefing.get("competitor", "").upper())
        set_run(p_label.add_run(label_left), size=13, bold=True, font=FONT_HEADING)
        set_run(p_label.add_run(L(lang, "cover_label_connector")),
                size=13, bold=True, color=CHERRY, font=FONT_HEADING)
        set_run(p_label.add_run(label_right), size=13, bold=True, font=FONT_HEADING)
    else:
        set_run(p_label.add_run(cover.get("label_left", "GO VOCAL")),
                size=13, bold=True, font=FONT_HEADING)

    for _ in range(1 if any_logo else 3):
        doc.add_paragraph()

    # Title (Chivo)
    p_title = doc.add_paragraph()
    set_run(p_title.add_run(cover.get("title", "")), size=46, font=FONT_HEADING)
    p_title.paragraph_format.line_spacing = 1.05
    p_title.paragraph_format.space_after = Pt(20)

    # Subtitle (Libre Franklin italic)
    if cover.get("subtitle"):
        p_sub = doc.add_paragraph()
        set_run(p_sub.add_run(cover["subtitle"]), size=13, italic=True, color=MEDIUM_PURPLE)
        p_sub.paragraph_format.line_spacing = 1.4
        p_sub.paragraph_format.space_after = Pt(18)

    page_break(doc)


def render_why_this_note(doc, briefing, lang):
    add_heading(doc, L(lang, "why_this_note"), level=2)
    for para in briefing.get("intro_paragraphs", []):
        add_body(doc, para, size=11)


def render_numbers(doc, briefing, lang):
    numbers = briefing.get("numbers", [])
    if not numbers:
        return
    add_heading(doc, L(lang, "go_vocal_in_numbers"), level=3)

    has_icons = any(_resolve(n.get("icon")) for n in numbers)
    n_rows = 3 if has_icons else 2
    table = doc.add_table(rows=n_rows, cols=len(numbers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    val_row = 1 if has_icons else 0

    for i, item in enumerate(numbers):
        if has_icons:
            icon_cell = table.cell(0, i)
            set_cell_bg(icon_cell, LILAC_1)
            icon_cell.text = ""
            p = icon_cell.paragraphs[0]; p.alignment = CENTER
            ic = _resolve(item.get("icon"))
            if ic:
                p.add_run().add_picture(str(ic), width=Cm(1.2))
            icon_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        top = table.cell(val_row, i)
        set_cell_bg(top, DARK_PURPLE_HEX)
        top.text = ""
        p = top.paragraphs[0]; p.alignment = CENTER
        set_run(p.add_run(item.get("value", "")), size=20, bold=True,
                color=WHITE, font=FONT_HEADING)
        top.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        bot = table.cell(val_row + 1, i)
        set_cell_bg(bot, LILAC_1)
        bot.text = ""
        p = bot.paragraphs[0]; p.alignment = CENTER
        set_run(p.add_run(item.get("label", "")), size=9)
        bot.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def render_usp_header(doc, number, title):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    for col, w in ((0, 2.5), (1, 14)):
        t.columns[col].width = Cm(w)
        t.rows[0].cells[col].width = Cm(w)
    num_cell = t.rows[0].cells[0]
    set_cell_bg(num_cell, DARK_PURPLE_HEX)
    num_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = num_cell.paragraphs[0]; p.alignment = CENTER
    set_run(p.add_run(str(number)), size=32, bold=True, color=WHITE, font=FONT_HEADING)
    title_cell = t.rows[0].cells[1]
    set_cell_bg(title_cell, LILAC_1)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_run(title_cell.paragraphs[0].add_run("  " + title), size=20, font=FONT_HEADING)
    doc.add_paragraph()


def _render_usp_text(container, usp, lang):
    p = container.add_paragraph()
    set_run(p.add_run(L(lang, "the_opportunity")), size=12, bold=True,
            color=CHERRY, font=FONT_HEADING)
    p.paragraph_format.space_after = Pt(2)
    add_body(container, usp.get("opportunity", ""), size=11)

    p = container.add_paragraph()
    set_run(p.add_run(L(lang, "what_we_offer")), size=12, bold=True,
            color=CHERRY, font=FONT_HEADING)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    if usp.get("features_intro"):
        add_body(container, usp["features_intro"], size=11)
    for feat in usp.get("features", []):
        add_bullet(container, feat.get("name", ""), feat.get("desc", ""))


def render_usp(doc, number, usp, lang):
    """Text first (full width), then the image below it — if one is provided.

    Image width: `width_cm` override, else 16cm for `placement: "full"` (wide
    shots), else 13cm centred. A missing image renders nothing (no placeholder).
    """
    render_usp_header(doc, number, usp.get("title", ""))
    _render_usp_text(doc, usp, lang)

    image = usp.get("image")
    if image:
        placement = image.get("placement", "")
        width = 16.0 if placement == "full" else 13.0
        add_image_slot(doc, image, width)

    doc.add_paragraph()


def render_usps(doc, briefing, lang):
    page_break(doc)
    add_heading(doc, L(lang, "where_we_stand_out"), level=2)
    if briefing.get("section_intro"):
        add_body(doc, briefing["section_intro"], size=11)
    for i, usp in enumerate(briefing.get("usps", []), start=1):
        render_usp(doc, i, usp, lang)


def render_bottom_line(doc, briefing, lang):
    bottom = briefing.get("bottom_line") or {}
    if not bottom:
        return
    page_break(doc)
    add_heading(doc, L(lang, "the_bottom_line"), level=2)
    for para in bottom.get("paragraphs", []):
        add_body(doc, para, size=11)
    if bottom.get("closer_italic"):
        add_body(doc, bottom["closer_italic"], size=11, italic=True, color=MEDIUM_PURPLE)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
def build(briefing: dict, output_path: Path) -> None:
    lang = briefing.get("language", "en")
    doc = setup_document()
    render_cover(doc, briefing, lang)
    render_why_this_note(doc, briefing, lang)
    render_numbers(doc, briefing, lang)
    render_usps(doc, briefing, lang)
    render_bottom_line(doc, briefing, lang)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main(argv: list[str]) -> int:
    global ASSETS_DIR
    if len(argv) not in (3, 4):
        sys.stderr.write("Usage: python generate_concept_note.py "
                         "<briefing.json> <output.docx> [assets_dir]\n")
        return 2
    briefing_path, output_path = Path(argv[1]), Path(argv[2])
    ASSETS_DIR = Path(argv[3]) if len(argv) == 4 else briefing_path.resolve().parent
    with briefing_path.open(encoding="utf-8") as f:
        briefing = json.load(f)
    build(briefing, output_path)
    print(f"Saved: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
